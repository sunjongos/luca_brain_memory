import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    if level == 0:
        run.font.color.rgb = RGBColor(44, 62, 80)
    else:
        run.font.color.rgb = RGBColor(41, 128, 185)

def add_paragraph(doc, text, bold_words=None):
    p = doc.add_paragraph()
    
    if bold_words:
        highlight_targets = [
            "남양주백병원 소화기센터 복강경 수술 클리닉 김용운 과장(외과 전문의, 아산병원 출신)",
            "김용운 과장님",
            "김용운 과장",
        ]
        
        current_text = text
        for target in highlight_targets:
            if target in current_text:
                parts = current_text.split(target)
                if len(parts) > 1:
                    p.add_run(parts[0])
                    run = p.add_run(target)
                    run.bold = True
                    run.font.color.rgb = RGBColor(22, 160, 133)
                    current_text = parts[1]
                break
                
        if current_text:
            if "패혈증" in current_text:
                parts = current_text.split("'패혈증'")
                if len(parts) > 1:
                    p.add_run(parts[0])
                    r = p.add_run("'패혈증'")
                    r.bold = True
                    r.font.color.rgb = RGBColor(217, 83, 79)
                    p.add_run(parts[1])
                    return p
            p.add_run(current_text)
        return p
    else:
        p.add_run(text)
    return p

document = Document()

title = document.add_heading('🏥 [건강정보] 담석 수술 고민 끝! 수많은 경험의 \'복강경 명의\' 김용운 과장이 팩트 체크 들어갑니다! 🩺', 0)
title_run = title.runs[0]
title_run.font.name = 'Malgun Gothic'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

subtitle_p = document.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_p.add_run('👨‍⚕️ 남양주백병원 소화기센터 복강경 수술 클리닉 김용운 과장(외과 전문의, 아산병원 출신)')
subtitle_run.bold = True
subtitle_run.font.color.rgb = RGBColor(243, 156, 18)
subtitle_run.font.name = 'Malgun Gothic'
subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

add_paragraph(document, "안녕하세요! 남양주백병원 소화기센터입니다. 😊", bold_words=True)
add_paragraph(document, "건강검진에서 우연히 발견된 '담석'! 당장 배가 아프지 않아서 수술을 미루고 계시나요?")
add_paragraph(document, "오늘은 남양주백병원 소화기센터의 든든한 대들보, 고난도 복강경 담낭 절제술의 스페셜리스트 남양주백병원 소화기센터 복강경 수술 클리닉 김용운 과장(외과 전문의, 아산병원 출신)님을 모시고 담석증 수술에 대한 모든 궁금증을 타파해 보겠습니다! 💡", bold_words=True)

add_heading(document, "💥 1. 안 아픈데 수술? \"패혈증 막으려면 골든타임 사수해야!\" 🚨", 2)
add_paragraph(document, "김용운 과장님은 외래 진료실에서 \"안 아프니까 그냥 두면 안 될까요?\"라고 묻는 환자분들께 늘 단호하게 말씀하십니다.", bold_words=True)
add_paragraph(document, "\"담석증을 가볍게 생각하고 방치하면 어느 순간 급성 담낭염으로 번집니다. 심해지면 담낭이 터지고 온몸에 염증이 퍼지는 '패혈증'이 올 수 있는 무서운 질환입니다.\"")
add_paragraph(document, "잦은 소화불량이 있거나 초음파상 염증이 보인다면, 생명을 위협하는 응급 상황이 오기 전에 실력 있는 든든한 아산병원 출신 전문의 김용운 과장님과 상담하여 예방적 수술을 결단하는 것이 가장 안전한 방법입니다! 🙌", bold_words=True)

add_heading(document, "✨ 2. '명의'의 손끝 + 초고속 회복(ERAS) = 일상으로 즉시 복귀! 🏃‍♂️", 2)
add_paragraph(document, "\"수술 후 통증이 무서워요 ㅠㅠ\"")
add_paragraph(document, "이런 걱정, 김용운 과장님 앞에서는 접어두셔도 좋습니다! 요즘 수술은 배에 구멍만 살짝 내는 '복강경 수술'로 진행되는데요. 특히 수많은 담석 환자를 치료해 온 김용운 과장님의 섬세하고 정교한 수술 테크닉은 이미 환자분들 사이에서 입소문이 자자합니다. 💪", bold_words=True)
add_paragraph(document, "저희 남양주백병원은 과장님의 명품 술기와 더불어 '수술 후 조기 회복 프로그램(ERAS)'을 완벽하게 적용 중입니다! 환자분들은 수술 후 통증을 최소화하면서도, 남들보다 훌쩍 빠른 속도로 건강하게 일상에 복귀하고 계십니다. 👍")

add_heading(document, "🤖 3. 대학병원 고가 로봇 수술? 김용운 과장의 '팩트 폭격!' 💣", 2)
add_paragraph(document, "대형 병원 갔다가 수천만 원을 호가하는 비싼 로봇 수술을 권유받고 놀란 가슴으로 찾아오시는 분들이 종종 계십니다. 이에 대한 김용운 과장님의 사이다 팩트 체크 들어갑니다! 👨‍⚕️✨", bold_words=True)
add_paragraph(document, "\"담낭 절제술의 글로벌 표준 치료는 로봇이 아닙니다. 경험 많은 의사가 다관절 복강경 수술 기구인 아티센셜(Artisential)와 세밀한 복강경으로 '직접' 뱃속 안의 구조물을 꼼꼼히 살피며 수술하는 것입니다.\"")
add_paragraph(document, "대학병원의 고가 로봇 수술은 가성비가 매우 떨어질 뿐만 아니라, 때로는 환자 치료 목적보다 병원 실적과 수익 창출을 위한 도구로 쓰이는 경우도 있다고 날카롭게 꼬집으셨습니다. 최상의 수술 결과는 '비싼 기계'가 아닌 명의의 풍부한 경험, 즉 '의사의 실력'에서 나옵니다! 💯")

add_heading(document, "🏥 '정직한 진료' 남양주백병원 소화기센터 복강경 수술 클리닉 김용운 과장", 2)
add_paragraph(document, "최상의 수술 결과를 위해 오직 환자분들의 무탈한 쾌유와 검증된 '표준 치료'만을 정직하게 고집하는 남양주백병원 소화기센터 복강경 수술 클리닉 김용운 과장(외과 전문의, 아산병원 출신)! 무리한 고가 수술을 권유하는 상술을 배제하고, 내 가족을 수술한다는 마음으로 가장 안전한 진료를 약속하십니다. 🤝", bold_words=True)
add_paragraph(document, "담석증 진단으로 두려우신가요? 더 이상 돌고 돌지 마시고, 풍부한 수술 노하우와 따뜻한 진료가 기다리는 김용운 과장님을 최우선으로 찾아오세요! 언제나 여러분의 든든한 해결사가 되어 드리겠습니다! 💪💚", bold_words=True)

try:
    img_path = r"C:\Users\USER\.gemini\antigravity\brain\1448649a-8b80-4ac3-89f4-28a4ffd9b85c\kim_yongwoon_anime_v2_1773807577340.png"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(img_path, width=Inches(5.0))
    
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run("수술 중인 김용운 과장의 친근한 애니메이션 아트워크")
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor(119, 119, 119)
    caption_run.font.name = 'Malgun Gothic'
    caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
except Exception as e:
    pass

for paragraph in document.paragraphs:
    for run in paragraph.runs:
        if run.font.name is None:
            run.font.name = 'Malgun Gothic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

output_path = r"C:\Users\USER\OneDrive\바탕 화면\언론기사\2026년 3월 언론기사\남양주백병원_담석증_블로그.docx"
document.save(output_path)
