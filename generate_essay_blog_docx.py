import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Batang'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Batang')
    if level == 0:
        run.font.color.rgb = RGBColor(30, 30, 30)
    else:
        run.font.color.rgb = RGBColor(50, 50, 50)

def add_paragraph(doc, text, is_quote=False, is_bold=False):
    p = doc.add_paragraph()
    
    if is_quote:
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        run = p.add_run(text)
        run.italic = False
        run.font.color.rgb = RGBColor(80, 80, 85)
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(50, 50, 50)
        run.font.size = Pt(11)
        if is_bold:
            run.bold = True

    run.font.name = 'Batang'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Batang')
    
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.6
    return p

document = Document()

category = document.add_paragraph()
category.alignment = WD_ALIGN_PARAGRAPH.CENTER
cat_run = category.add_run("솔직한 일상 기록")
cat_run.font.size = Pt(9)
cat_run.font.color.rgb = RGBColor(120, 120, 120)
cat_run.font.name = 'Batang'
cat_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Batang')

title = document.add_heading('로봇 수술의 권유 앞에서, 진짜 의사를 만나다', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_paragraph(document, "매년 받는 직장인 건강검진. 늘 그렇듯 '운동 부족'이나 '경도 비만' 정도를 예상하며 무심코 넘기던 결과지에서 낯선 단어를 발견했다. '담석'. 당장 증상이 있는 건 아니었지만, 몸속에 돌이 자라고 있다는 사실은 일상의 묘한 불청객처럼 신경이 쓰였다.")

add_paragraph(document, "시간이 지날수록 소화가 잘 안 되고 이따금 오른쪽 윗배가 뻐근해지기 시작했다. 인터넷을 찾아보니 담석을 방치하다 급성 담낭염으로 번져 응급실에 실려 가는 사례가 수두룩했다. 결국 수술을 결심하고 이름만 대면 아는 서울의 대형 병원에 예약을 잡았다.")

add_paragraph(document, "대기실은 시장통처럼 붐볐고, 진료실에서 의사와 대면한 시간은 채 3분도 되지 않았다. 모니터를 보며 빠르게 차트를 넘기던 의사는 아주 당연하다는 듯이 '로봇 수술'을 제안했다. \"초정밀 수술이라 회복이 빠릅니다.\" 설명은 그게 다였다. 하지만 수납처에서 확인한 수술비는 수천만 원에 달했다. 내 생명이 달린 일이라지만, 뭔가 단단히 잘못 돌아가고 있다는 씁쓸함이 밀려왔다. 담낭 하나 떼어내는데, 정말 이 거대한 비용과 화려한 기계만이 유일한 정답일까?")

add_paragraph(document, "확신이 서지 않던 차에, 지인의 소개로 남양주백병원 소화기센터를 찾았다. 복강경 수술 클리닉의 김용운 과장님(외과 전문의, 아산병원 출신)과의 첫 진료. 그곳의 공기는 대형 병원의 컨베이어 벨트 같은 서늘함과는 사뭇 달랐다.")

add_paragraph(document, "그분은 내 초음파 사진을 하나하나 짚어가며 현재 상태를 설명해 주셨다. 조심스레 로봇 수술에 대해 묻는 내게, 과장님은 잔잔하지만 단단한 어조로 말씀하셨다.")

add_paragraph(document, "\"수천만 원에 달하는 로봇 수술, 굳이 하지 않으셔도 됩니다. 담낭 절제술의 가장 검증된 표준 치료는 여전히 복강경 수술입니다. 통원 치료가 필요할 만큼 위중하거나 섬세한 암 조직 절제가 아니라면, 기계의 도움이 반드시 필요한 것은 아닙니다. 제 손끝의 감각, 그리고 다관절 복강경 수술 기구인 '아티센셜(Artisential)'이면 충분합니다. 제 손으로 꼼꼼히 살피고 정교하게 매듭지어 안전하게 마무리해 드리겠습니다.\"", True)

add_paragraph(document, "그제야 굳어있던 어깨에 힘이 풀렸다. 무언가에 쫓기듯 수술대에 오르지 않아도 된다는 안도감. 불필요한 과잉 진료 대신, 환자에게 가장 합리적이고 안전한 길을 제시하는 진짜 전문가의 모습이 거기 있었다.")

add_paragraph(document, "수술 당일의 기억은 흐릿하지만, 며칠간 앓아누울 거란 걱정은 기우에 불과했다. 남양주백병원의 '수술 후 조기 회복 프로그램(ERAS)' 덕분인지 마취에서 깬 후 통증은 생각보다 가벼웠고, 난 다음 날 바로 침대에서 내려와 복도를 걸었다. 과장님의 섬세한 술기가 몸에 무리를 주지 않았다는 가장 확실한 증거였다.")

add_paragraph(document, "지금 배에는 아주 자그마한, 거의 눈에 띄지 않는 흉터 두세 개가 전부다. 화려한 대학병원의 로봇 수술을 받지 않은 것을 단 한 번도 후회한 적이 없다. 오히려, 상술과 실적의 압박 속에서도 흔들림 없이 '외과 의사의 본질'을 지켜가고 있는 정직한 명의를 만난 것이 내겐 더없는 행운이었다.")

add_paragraph(document, "몸속의 돌만 빼낸 것이 아니라, 차가웠던 병원에 대한 불신 한 조각까지 떼어내 준 김용운 과장님께 잔잔한 감사의 인사를 보낸다.")

try:
    img_path = r"C:\Users\USER\.gemini\antigravity\brain\1448649a-8b80-4ac3-89f4-28a4ffd9b85c\kim_yongwoon_anime_v2_1773807577340.png"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(img_path, width=Inches(4.5))
    
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run("과잉 진료 대신 치료의 본질을 짚어주던, 든든했던 김용운 과장님의 옅은 미소.")
    caption_run.italic = False
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor(120, 120, 120)
    caption_run.font.name = 'Batang'
    caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Batang')
except Exception as e:
    pass

output_path = r"C:\Users\USER\OneDrive\바탕 화면\언론기사\2026년 3월 언론기사\남양주백병원_담석증_블로그_수필형.docx"
document.save(output_path)
