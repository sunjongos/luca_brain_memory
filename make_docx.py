import docx
import os

md_path = r"C:\Users\USER\OneDrive\바탕 화면\뇌신경센터\sleep_apnea_stroke_article.md"
docx_path = r"C:\Users\USER\OneDrive\바탕 화면\언론기사\2026년 3월 언론기사\뇌신경센터\sleep_apnea_stroke_article.docx"

doc = docx.Document()

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

for line in lines:
    line = line.strip()
    if not line:
        continue
    if line.startswith('#'):
        heading_level = line.count('#', 0, 6)
        text = line.replace('#', '').strip()
        doc.add_heading(text, level=heading_level if heading_level <= 9 else 1)
    else:
        # Strip markdown bolding for simplicity or just add it
        text = line.replace('**', '')
        doc.add_paragraph(text)

os.makedirs(os.path.dirname(docx_path), exist_ok=True)
doc.save(docx_path)
print("DOCX created successfully at:", docx_path)
