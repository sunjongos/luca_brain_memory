import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    if level == 0:
        run.font.color.rgb = RGBColor(0, 86, 179)
    else:
        run.font.color.rgb = RGBColor(0, 64, 133)

def add_paragraph(doc, text, bold_words=None):
    p = doc.add_paragraph()
    
    if bold_words:
        highlight_targets = [
            "남양주백병원 소화기센터 복강경 수술 클리닉 김용운 과장(외과 전문의, 아산병원 출신)",
            "김용운 과장",
            "김용운 과장 수술팀",
            "남양주백병원 소화기센터 복강경 수술 클리닉의 김용운 과장 수술팀",
            "남양주백병원 소화기센터 복강경 수술 클리닉 김용운 과장"
        ]
        
        current_text = text
        for target in highlight_targets:
            if target in current_text:
                parts = current_text.split(target)
                if len(parts) > 1:
                    p.add_run(parts[0])
                    run = p.add_run(target)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 86, 179)
                    current_text = parts[1]
                break
                
        if current_text:
            if "패혈증" in current_text:
                parts = current_text.split("패혈증")
                if len(parts) > 1:
                    p.add_run(parts[0])
                    r = p.add_run("패혈증")
                    r.bold = True
                    r.font.color.rgb = RGBColor(217, 83, 79)
                    current_text = parts[1]
            p.add_run(current_text)
        return p
    else:
        p.add_run(text)
    return p

document = Document()

title = document.add_heading('[건강칼럼] 검진에서 발견된 담석, 수술 고민 끝! \'복강경 수술 권위자\' 김용운 과장이 답하다', 0)
title_run = title.runs[0]
title_run.font.name = 'Malgun Gothic'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

add_paragraph(document, "최근 건강검진을 받다가 우연히 담석을 발견하고 병원을 찾는 환자들이 늘고 있다. 증상이 없는 경우 당장 수술해야 할지 고민하는 경우가 많다. 이와 관련해 남양주백병원 소화기센터 복강경 수술 클리닉 김용운 과장(외과 전문의, 아산병원 출신)은 \"방치할 경우 심각한 합병증을 유발할 수 있어 정밀 진단과 선제적 치료가 필수\"라고 강조했다.", bold_words=True)

add_heading(document, "담낭 절제술, 안 아프다고 방치하면 '패혈증' 시한폭탄", 2)
add_paragraph(document, "담석증을 가볍게 여겨 치료 시기를 놓치면 급성 담낭염으로 발전할 수 있다. 김용운 과장은 \"담낭염이 악화되면 담낭이 터지거나 염증이 전신으로 퍼져 패혈증을 유발하며, 이는 생명을 위협할 수 있는 초응급 상태\"라고 경고하며, \"담석으로 인한 잦은 소화불량이나 우상복부 통증, 염증 소견이 보인다면 예방적 차원에서의 담낭 절제술이 강력히 권고된다\"고 설명했다.", bold_words=True)

add_heading(document, "'단기 회복'의 마술, 아산병원 출신 김용운 과장의 복강경 수술 & ERAS 프로그램", 2)
add_paragraph(document, "최근의 담낭 절제술은 대부분 절개 부위를 최소화하는 복강경 수술로 진행된다. 특히 남양주백병원 소화기센터 복강경 수술 클리닉의 김용운 과장 수술팀은 이 복강경 수술의 장점을 극대화하여 ‘수술 후 조기 회복 프로그램(ERAS, Enhanced Recovery After Surgery)’을 선도적으로 적용하고 있다. 김용운 과장의 풍부한 임상 경험과 섬세한 술기 덕분에 환자들의 수술 후 통증은 획기적으로 줄고, 당일 혹은 익일 바로 일상 복귀가 가능할 정도로 뛰어난 예후를 보이고 있다.", bold_words=True)

add_heading(document, "상술로 변질된 대학병원 로봇 수술... \"최고의 수술 도구는 집도의의 검증된 손끝\"", 2)
add_paragraph(document, "최근 일부 대형 병원에서 다양한 질환에 무분별하게 수천만 원에 달하는 고가의 로봇 수술을 권유하며 환자들의 경제적 부담을 가중시키고 있다.")
add_paragraph(document, "김용운 과장은 이에 대해 쓴소리를 아끼지 않았다. \"일부 전립선암 수술 등 세밀한 공간 조작이 절대적인 특수 질환을 제외하면, 담낭 절제술의 글로벌 표준 치료는 숙련된 전문의가 복강경과 다관절 복강경 수술 기구인 아티센셜(Artisential)를 이용해 '직접' 담낭을 절제하는 것입니다.\"", bold_words=True)
add_paragraph(document, "그는 또한 \"대학병원에서 담낭 절제술에까지 굳이 고비용의 비급여 로봇 수술을 강요하는 것은 환자의 치료적 이득보다는 기관의 로봇 수술 실적 건수 채우기나 병원 수익 창출 목적일 가능성을 배제할 수 없다\"고 일침을 가했다. 수술의 완벽한 결과를 좌우하는 것은 값비싼 기계 장비가 아니라, 오직 의사의 풍부한 임상 경험과 정교한 직접 술기라는 것이다.", bold_words=True)
add_paragraph(document, "남양주백병원 소화기센터 복강경 수술 클리닉 김용운 과장(외과 전문의, 아산병원 출신)은 불필요한 시술이나 과도한 비급여 항목을 배제하고, 가장 안전하며 검증된 '표준 치료'만을 고집하는 정직한 의술로 지역민들의 깊은 신뢰를 받고 있다. 건강검진에서 담석 진단을 받았다면, 로봇 수술의 환상에 기대기보다 실력이 검증된 전문의를 찾아 최적의 치료 시기를 놓치지 않는 것이 현명한 선택이다.", bold_words=True)

try:
    img_path = r"C:\Users\USER\.gemini\antigravity\brain\1448649a-8b80-4ac3-89f4-28a4ffd9b85c\kim_yongwoon_real_v2_1773807561722.png"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(img_path, width=Inches(5.0))
    
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run("(최고의 술기로 수술을 리드하는 남양주백병원 김용운 과장 수술 현장)")
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor(119, 119, 119)
    caption_run.font.name = 'Malgun Gothic'
    caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
except Exception as e:
    pass

for paragraph in document.paragraphs:
    for run in paragraph.runs:
        if run.font.name is None:
            run.font.name = 'Malgun Gothic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

output_path = r"C:\Users\USER\OneDrive\바탕 화면\언론기사\2026년 3월 언론기사\남양주백병원_담석증_기사.docx"
document.save(output_path)
