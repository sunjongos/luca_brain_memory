import docx
from docx.shared import Inches
from PIL import Image
import os

md_path = r"C:\Users\USER\OneDrive\바탕 화면\뇌신경센터\sleep_apnea_stroke_article.md"
img_path = r"C:\Users\USER\OneDrive\바탕 화면\언론기사\2026년 3월 언론기사\뇌신경센터\sleep_apnea_stroke_image.png"
jpg_img_path = img_path.replace(".png", "_conv.jpg")
docx_path = r"C:\Users\USER\OneDrive\바탕 화면\언론기사\2026년 3월 언론기사\뇌신경센터\sleep_apnea_stroke_article.docx"

# Convert image to JPG to be safe for python-docx
try:
    with Image.open(img_path) as img:
        img.convert("RGB").save(jpg_img_path, "JPEG")
except Exception as e:
    print(f"Image conversion error: {e}")
    # Fallback to original
    jpg_img_path = img_path

doc = docx.Document()

# Add picture
try:
    doc.add_picture(jpg_img_path, width=Inches(6.0))
except Exception as e:
    print(f"Error adding picture to DOCX: {e}")

# Add content
try:
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('#'):
            lvl = line.count('#', 0, 6)
            doc.add_heading(line.replace('#', '').strip(), level=lvl if lvl <= 9 else 1)
        else:
            doc.add_paragraph(line.replace('**', ''))
except Exception as e:
    print(f"Error reading MD: {e}")

doc.save(docx_path)
print("DOCX with image generated successfully at:", docx_path)
