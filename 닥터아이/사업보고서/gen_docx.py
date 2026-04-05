from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

REPORT_PATH = r'C:\Users\USER\OneDrive\바탕 화면\luca연구에이전트\닥터아이\사업보고서\사업보고서_닥터아이_MultiAgent.docx'

doc = Document()

# ── 페이지 여백 설정
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

def set_font(run, size=11, bold=False, color=None):
    run.font.name = '나눔고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    colors = {1:(11,31,58), 2:(26,77,179), 3:(0,150,130)}
    sizes  = {1:18, 2:14, 3:12}
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, sizes.get(level,12), bold=True, color=colors.get(level,(0,0,0)))
    return p

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, 10)
    return p

def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, (cell, data) in enumerate(zip(row.cells, cells_data)):
        cell.text = data
        run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(data)
        run.font.name = '나눔고딕'
        run.font.size = Pt(9)
        run.font.bold = header
        if header:
            run.font.color.rgb = RGBColor(255,255,255)
            # 배경색
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '0B1F3A')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)

# ===================== COVER =====================
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run('\n\n닥터아이 Multi-Agent\n당뇨 눈 질환 관리 시스템 사업보고서\n\n')
set_font(r, 22, bold=True, color=(11,31,58))

cover2 = doc.add_paragraph()
cover2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = cover2.add_run('세계 최초 — AI 안저진단 + 다중 에이전트 오케스트레이션으로\n당뇨성 실명을 막는 의료 자동화 플랫폼 사업 제안\n\n')
set_font(r2, 12, color=(70,90,130))

cover3 = doc.add_paragraph()
cover3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = cover3.add_run('PaiHealthcare-LCK 연구소 | 닥터아이 | 2026년 03월 07일 | 대외비\n\n')
set_font(r3, 10, color=(100,120,160))

doc.add_page_break()

# ===================== 1. 사업 개요 =====================
add_heading('1. 사업 개요', 1)
add_body('본 사업은 파이헬스케어-LCK 연구소가 개발한 AI 안저진단 솔루션 닥터아이(Doctor Eye)를 핵심 엔진으로, Multi-Agent AI 시스템을 결합하여 당뇨 환자의 눈 질환을 자동화된 3단계 워크플로우로 관리하는 세계 최초의 의료 AI 오케스트레이션 플랫폼을 구축하는 것을 목표로 합니다.')

stats = [
    ('구분', '수치', '설명'),
    ('AI 진단 대상', '4대 질환', '당뇨망막증·녹내장·황반변성·망막혈관폐쇄'),
    ('전문 AI 에이전트', '7개 이상', '병렬·직렬 실행으로 워크플로우 자동화'),
    ('조기 발견 예방률', '95%', '당뇨 합병증 실명 예방 의학적 근거'),
    ('AI 판독 시간', '5분 이내', '촬영 후 전자동 리포트 생성 및 EMR 연동'),
]
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_row(t, stats[0], header=True)
for row_data in stats[1:]:
    add_table_row(t, row_data)
doc.add_paragraph()

add_body('핵심 혁신: 기존 AI 의료 솔루션이 진단에 그쳤다면, 본 시스템은 진단→처방 초안→협진 의뢰→환자 알림→재방문 예약까지 의료 업무 전 과정을 Multi-Agent가 자동 처리하며, 중요한 의사결정만 의사가 Human-in-the-Loop 방식으로 승인합니다.')

doc.add_page_break()

# ===================== 2. 문제 정의 =====================
add_heading('2. 문제 정의 — 왜 지금인가', 1)
add_heading('2.1 당뇨 합병증 실명의 위험성', 2)
add_body('대한민국 당뇨 환자는 약 600만 명으로, 이 중 당뇨성 망막증 발생률은 진단 10년 후 30%에 달합니다. 당뇨 합병증 중 실명은 삶의 질을 가장 극단적으로 파괴합니다. 당뇨성 실명의 95%는 조기 발견 시 예방 가능합니다.')

add_heading('2.2 현재 의료 시스템의 한계', 2)
prob = [
    ('문제', '현황', '결과'),
    ('안저검사 누락', '당뇨 외래 안저검사 시행률 30% 미만', '진단 지연 → 실명'),
    ('단절된 워크플로우', '진단·처방·협진이 각각 분리 운영', '의료 질 편차 발생'),
    ('의사 업무 과부하', '반복적 오더 작성·협진 의뢰 수작업', '진료 집중도 저하'),
    ('환자 추적 실패', '재방문 주기 관리 없음', '고위험군 이탈'),
]
t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
add_table_row(t2, prob[0], header=True)
for row_data in prob[1:]:
    add_table_row(t2, row_data)
doc.add_paragraph()

doc.add_page_break()

# ===================== 3. 솔루션 =====================
add_heading('3. 솔루션 — 닥터아이 + Multi-Agent 시스템', 1)
add_heading('3.1 시스템 아키텍처 (4개 레이어)', 2)

layers = [
    ('Layer 1 — 진단 레이어 (닥터아이 AI)', '안저 촬영 → 4대 안질환 AI 분석 → 위험도 스코어링 → 결과 분기 판단 (완전 자동)'),
    ('Layer 2 — 처리 레이어 (Multi-Agent 병렬 실행)', 'Prescription Agent(처방 초안) + ECAS Agent(혈관 평가) + Notify Agent(환자 문자) + Scheduler Agent(재방문 예약) — 동시 병렬 실행'),
    ('Layer 3 — 승인 레이어 (Human-in-the-Loop)', '처방 최종 승인 · 협진 의뢰 서명 · 고위험군 뇌 MRI 의뢰 — 반드시 의사 개입'),
    ('Layer 4 — 학습 레이어 (Feedback Loop)', '의사 수정 이력 학습 → 처방 초안 정확도 지속 향상 → 개인화 프로토콜 진화'),
]
for title, desc in layers:
    p = doc.add_paragraph(style='List Bullet')
    r_title = p.add_run(title + ': ')
    set_font(r_title, 10, bold=True, color=(26,77,179))
    r_desc = p.add_run(desc)
    set_font(r_desc, 10)

add_heading('3.2 Agent 구성 및 역할', 2)
agents = [
    ('Agent명', '역할', '실행방식', '의사개입'),
    ('Screening Agent', '닥터아이 결과 수신·위험도 분류', '자동', '없음'),
    ('Risk Scoring Agent', 'X-ECAS 경동맥 위험도 계산', '자동', '없음'),
    ('Prescription Draft Agent', 'EMR 오더셋 처방 초안 생성', '자동 초안', '승인 필요'),
    ('Referral Agent', '협진 의뢰서 자동 작성', '자동 초안', '서명 필요'),
    ('Patient Notify Agent', '결과 문자·앱 알림 발송', '자동', '없음'),
    ('Scheduler Agent', '다음 방문 일정 자동 예약', '자동', '없음'),
    ('Alert Agent', 'High-Risk 긴급 의사 알림', '즉시 알림', '필수'),
    ('Learning Agent', '처방 이력 학습·모델 개선', '배치 자동', '없음'),
]
t3 = doc.add_table(rows=1, cols=4)
t3.style = 'Table Grid'
add_table_row(t3, agents[0], header=True)
for row_data in agents[1:]:
    add_table_row(t3, row_data)

doc.add_page_break()

# ===================== 4. 3단계 워크플로우 =====================
add_heading('4. 3단계 임상 워크플로우', 1)
add_heading('Stage 1 — 닥터아이 AI 안저검사', 2)
add_body('당뇨 외래 내원 시 E6670 코드로 안저촬영 시행. 닥터아이 AI가 5분 내 4대 안질환 분석 및 위험도 리포트 자동 생성 → EMR 팝업 연동. 정상 판정 시 6~12개월 후 자동 재방문 예약. 이상 소견 시 2단계로 자동 이행.')
add_heading('Stage 2 — 내과 조치 및 약물 처방 (이상 소견 시)', 2)
add_body('Prescription Draft Agent가 표준 프로토콜 기반 처방 초안 생성 (칼슘 도베실레이트 bid + 징코빌로바 추출물). 의사가 검토·수정 후 최종 승인. 동시에 X-ECAS 검사 오더 자동 생성.')
add_heading('Stage 3 — X-ECAS 위험도별 처치', 2)
risk = [
    ('위험도', '약물', '검사', '협진'),
    ('LOW', '기존 요법 유지', '닥터아이 6~12개월 추적', '불필요'),
    ('INTERMEDIATE', '스타틴 + 항혈소판제', '경동맥 재초음파 6~12개월', '순환기내과 컨설트'),
    ('HIGH', '고강도 스타틴 + 항혈소판제', '뇌 MRI+MRA (의사 필수 승인)', '신경과·순환기내과 협진'),
]
t4 = doc.add_table(rows=1, cols=4)
t4.style = 'Table Grid'
add_table_row(t4, risk[0], header=True)
for row_data in risk[1:]:
    add_table_row(t4, row_data)

doc.add_page_break()

# ===================== 5. 세계 최초 포지셔닝 =====================
add_heading('5. 세계 최초 포지셔닝 및 경쟁 우위', 1)
comp = [
    ('비교 항목', '기존 의료 AI', '닥터아이 Multi-Agent'),
    ('AI 역할', '단일 진단만', '진단→처방→협진→알림 자동화'),
    ('워크플로우', '의사가 모든 단계 수동 처리', 'Multi-Agent 병렬·직렬 자동화'),
    ('의사 역할', '모든 결정', '핵심 결정만 Human-in-the-Loop'),
    ('학습 능력', '정적 모델', '처방 이력으로 지속 개인화 학습'),
    ('EMR 연동', '별도 입력 필요', '위험도 기반 오더셋 자동 팝업'),
    ('환자 관리', '의사 기억·수동', 'Scheduler Agent 자동 추적 관리'),
]
t5 = doc.add_table(rows=1, cols=3)
t5.style = 'Table Grid'
add_table_row(t5, comp[0], header=True)
for row_data in comp[1:]:
    add_table_row(t5, row_data)
doc.add_paragraph()
add_body('2026년 현재, 안과 영역(특히 당뇨 눈 질환)에 Multi-Agent 오케스트레이션 + Human-in-the-Loop를 결합한 임상 워크플로우 시스템은 국내외 논문 및 상용 제품에서 발표된 바 없음. 세계 최초 출시 가능 영역.')

doc.add_page_break()

# ===================== 6. 시장 분석 =====================
add_heading('6. 시장 분석', 1)
mkt = [
    ('구분', '규모', '비고'),
    ('글로벌 안과 AI 시장 (2028)', '$8.7B', '연평균 성장률 38.7%'),
    ('국내 당뇨 환자', '600만명', '고령화로 2030년 700만명 이상 예상'),
    ('국내 내과·가정의학과 의원', '3만+', '1차 타깃: 당뇨 전문 클리닉 보유 병원'),
    ('글로벌 의료 AI 자동화 시장', '$42B', 'Multi-Agent 의료 워크플로우 급성장'),
]
t6 = doc.add_table(rows=1, cols=3)
t6.style = 'Table Grid'
add_table_row(t6, mkt[0], header=True)
for row_data in mkt[1:]:
    add_table_row(t6, row_data)
doc.add_paragraph()

add_heading('6.1 수익 구조 예상', 2)
rev = [
    ('시기', '도입 병원', '월 매출 추정', '핵심 활동'),
    ('2026 (파일럿)', '5개', '500만원', '임상 검증, 데이터 수집'),
    ('2027 (초기 확장)', '50개', '5,000만원', 'EMR 연동 완성, 학술 발표'),
    ('2028 (성장)', '300개', '3억원', '전국 확산, 보험 수가 등재'),
    ('2029 (글로벌)', '500개+', '10억원+', '동남아·중동 수출'),
]
t7 = doc.add_table(rows=1, cols=4)
t7.style = 'Table Grid'
add_table_row(t7, rev[0], header=True)
for row_data in rev[1:]:
    add_table_row(t7, row_data)

doc.add_page_break()

# ===================== 7. 특허 전략 =====================
add_heading('7. 특허 및 지식재산 전략', 1)
patents = [
    ('특허 ①', '당뇨 눈 질환 Multi-Agent 자동화 진료 시스템', 'AI 안저진단 결과를 트리거로 복수 AI 에이전트가 병렬·직렬로 처방 초안·협진·알림·일정을 자동 처리하는 의료 워크플로우 시스템 및 방법'),
    ('특허 ②', 'Human-in-the-Loop 기반 의료 AI 처방 승인 방법', 'AI 에이전트가 생성한 처방 초안에 대해 의사가 단계적으로 검토·수정·승인하는 구조의 의료 의사결정 지원 시스템'),
    ('특허 ③', '안저·경동맥 통합 위험도 AI 스코어링 시스템', 'X-eye 안저 AI 판독 결과와 X-ECAS 경동맥 초음파 데이터를 AI로 통합 분석하여 뇌졸중·심근경색 위험도를 복합 산출하는 방법'),
    ('특허 ④', '의사 처방 이력 학습 개인화 EMR 오더셋 추천 방법', '각 의사의 처방 패턴을 학습하여 맞춤형 오더셋을 자동 추천하는 적응형 AI 시스템'),
]
for idx, title, desc in patents:
    p = doc.add_paragraph(style='List Bullet')
    r_t = p.add_run(f'{idx} {title}: ')
    set_font(r_t, 10, bold=True, color=(124,58,237))
    r_d = p.add_run(desc)
    set_font(r_d, 10)

doc.add_page_break()

# ===================== 8. 실행 계획 =====================
add_heading('8. 실행 계획 (로드맵)', 1)
road = [
    ('시기', '단계', '주요 목표'),
    ('2026 Q2', '파일럿', '닥터아이 + 2개 Agent 연동, 파일럿 병원 3~5곳, IRB 승인'),
    ('2026 Q4', '시스템 완성', '8개 Agent 전체 구축, EMR 3종 연동, 특허 출원 4건'),
    ('2027', '초기 확산', '50개 병원 도입, 보험 수가 등재 신청, Series A 투자 유치'),
    ('2028', '전국 확장', '300개 병원 목표, 건강보험 연동, 동남아 파트너 계약'),
    ('2029+', '글로벌', '해외 수출 본격화, 당뇨 외 질환 확장, IPO 준비'),
]
t8 = doc.add_table(rows=1, cols=3)
t8.style = 'Table Grid'
add_table_row(t8, road[0], header=True)
for row_data in road[1:]:
    add_table_row(t8, row_data)

doc.add_page_break()

# ===================== 9. 결론 =====================
add_heading('9. 결론 — 지금이 기회입니다', 1)
add_body('당뇨성 실명은 막을 수 있는 실명입니다. 닥터아이가 보고, Multi-Agent가 움직이고, 의사가 결정합니다. 기술적 혁신 + 의학적 근거 + 세계 최초 포지셔닝이 갖춰진 이 시스템은 대한민국을 의료 AI 자동화의 글로벌 리더로 만들 수 있습니다.')
add_body('본 사업에 관심 있는 투자자, 병원 파트너, 기술 협력사는 PaiHealthcare-LCK 연구소로 문의 바랍니다.')

doc.save(REPORT_PATH)
print(f'DOCX 저장 완료: {REPORT_PATH}')
