import os
import markdown
import docx
from docx.shared import Pt
import traceback

def create_files():
    md_path = r"c:\Users\sunjo\Desktop\ontology\DoctorEye_MultiModal_RND_Report.md"
    html_path = r"c:\Users\sunjo\Desktop\ontology\DoctorEye_MultiModal_RND_Report.html"
    docx_path = r"c:\Users\sunjo\Desktop\ontology\DoctorEye_MultiModal_RND_Report.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: Could not find {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # 1. HTML Dashboard Generation
    html_content = markdown.markdown(md_text, extensions=['fenced_code', 'tables'])
    html_template = f"""
    <!DOCTYPE html>
    <html lang="ko">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Doctor Eye MultiModal R&D Report</title>
        <style>
            body {{ font-family: 'Pretendard', sans-serif; line-height: 1.6; color: #333; max-width: 900px; margin: 0 auto; padding: 20px; background-color: #f7f9fc; }}
            .dashboard {{ background: white; padding: 40px; border-radius: 12px; box-shadow: 0 10px 30px rgba(0,0,0,0.1); }}
            h1, h2, h3 {{ color: #2c3e50; border-bottom: 2px solid #ecf0f1; padding-bottom: 10px; }}
            h1 {{ color: #2980b9; }}
            pre {{ background: #2c3e50; color: #ecf0f1; padding: 15px; border-radius: 8px; overflow-x: auto; }}
            a {{ color: #3498db; text-decoration: none; }}
            a:hover {{ text-decoration: underline; }}
            ul, ol {{ margin-left: 20px; }}
            li {{ margin-bottom: 8px; }}
        </style>
        <script type="module">
            import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
            mermaid.initialize({{ startOnLoad: true, theme: 'default' }});
        </script>
    </head>
    <body>
        <div class="dashboard">
            {html_content.replace('class="language-mermaid"', 'class="mermaid"')}
        </div>
    </body>
    </html>
    """
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_template)
    print(f"✅ HTML Dashboard generated: {html_path}")

    # 2. DOCX Generation
    try:
        doc = docx.Document()
        doc.add_heading('닥터아이(Doctor Eye) 멀티모달 R&D 기획 보고서', 0)
        
        for line in md_text.split('\n'):
            line = line.strip()
            if not line: continue
            if line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. '):
                doc.add_paragraph(line[3:], style='List Number')
            else:
                doc.add_paragraph(line)
                
        doc.save(docx_path)
        print(f"✅ DOCX Document generated: {docx_path}")
    except Exception as e:
        print(f"❌ Failed to generate DOCX: {e}")
        traceback.print_exc()

if __name__ == "__main__":
    create_files()
