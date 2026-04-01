import sys
import subprocess

def install_docx():
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

install_docx()

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    if level == 0:
        run.font.color.rgb = RGBColor(0, 86, 179)
    else:
        run.font.color.rgb = RGBColor(0, 64, 133)

def add_paragraph(doc, text, bold_words=None):
    p = doc.add_paragraph()
    
    if bold_words:
        parts = text.split("패혈증")
        if len(parts) > 1:
            p.add_run(parts[0])
            run = p.add_run("패혈증")
            run.bold = True
            run.font.color.rgb = RGBColor(217, 83, 79)
            p.add_run(parts[1])
            return p
            
        parts = text.split("복강경 수술")
        if len(parts) > 1 and "진행된다" in text:
            p.add_run(parts[0])
            run = p.add_run("복강경 수술")
            run.bold = True
            p.add_run(parts[1])
            return p
            
        parts = text.split("수술 후 조기 회복 프로그램(ERAS, Enhanced Recovery After Surgery)")
        if len(parts) > 1:
            p.add_run(parts[0])
            run = p.add_run("수술 후 조기 회복 프로그램(ERAS, Enhanced Recovery After Surgery)")
            run.bold = True
            p.add_run(parts[1])
            return p
            
        parts = text.split("표준 치료는 숙련된 전문의가 복강경과 첨단 필수 수술 기구(Advanced essentials tools)를 이용하여 직접 담낭을 절제하는 것")
        if len(parts) > 1:
            p.add_run(parts[0])
            run = p.add_run("표준 치료는 숙련된 전문의가 복강경과 첨단 필수 수술 기구(Advanced essentials tools)를 이용하여 직접 담낭을 절제하는 것")
            run.bold = True
            p.add_run(parts[1])
            return p
            
        parts = text.split("의사의 풍부한 임상 경험과 정교한 직접 술기")
        if len(parts) > 1:
            p.add_run(parts[0])
            run = p.add_run("의사의 풍부한 임상 경험과 정교한 직접 술기")
            run.bold = True
            run.font.color.rgb = RGBColor(217, 83, 79)
            p.add_run(parts[1])
            return p

    p.add_run(text)
    return p

document = Document()

# Title
title = document.add_heading('[건강칼럼] 검진에서 발견된 담석, 수술이 꼭 필요한 경우는?', 0)
title_run = title.runs[0]
title_run.font.name = 'Malgun Gothic'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# Author
author_p = document.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
author_run = author_p.add_run('남양주백병원 소화기센터 일반외과 전문의 김요운 과장')
author_run.bold = True
author_run.font.name = 'Malgun Gothic'
author_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# Intro
add_paragraph(document, "최근 건강검진을 받다가 우연히 담석을 발견하고 병원을 찾는 환자들이 늘고 있다. 증상이 없는 경우 당장 수술을 해야 할지 고민하는 경우가 많다. 하지만 방치할 경우 심각한 합병증을 유발할 수 있어 정확한 진단과 치료가 필수적이다.")

# Sec 1
add_heading(document, "담낭 절제술, 왜 꼭 해야 하는가?", 2)
add_paragraph(document, "담석증을 가볍게 여겨 치료 시기를 놓치면 급성 담낭염으로 발전할 수 있다. 담낭염이 심해지면 담낭이 터지거나 염증이 전신으로 퍼져 패혈증을 유발하며, 이는 생명을 위협할 수 있는 매우 위중한 상태가 된다. 따라서 담석으로 인한 통증이나 염증 소견이 보인다면 예방적 차원에서라도 담낭 절제술을 적극적으로 고려해야 한다.", bold_words=True)

# Sec 2
add_heading(document, "복강경 수술과 수술 후 조기 회복 프로그램(ERAS)", 2)
add_paragraph(document, "과거와 달리 최근의 담낭 절제술은 대부분 복강경 수술로 진행된다. 절개 부위가 작아 통증이 적고 회복이 빠르다는 장점이 있다. 특히 남양주백병원에서는 복강경 수술의 장점을 극대화하여 수술 후 조기 회복 프로그램(ERAS, Enhanced Recovery After Surgery)을 적극 적용하고 있다. 이를 통해 환자는 수술 후 통증을 최소화하고 가장 빠른 일상 복귀가 가능하다.", bold_words=True)

# Sec 3
add_heading(document, "로봇 수술의 빛과 그림자... 담낭 절제술의 '표준 치료'는?", 2)
add_paragraph(document, "최근 대형 대학병원 등에서 다양한 질환에 경쟁적으로 로봇 수술을 도입하며, 담낭 절제술에도 이를 권유하는 경우가 많다. 그러나 환자들은 로봇 수술 권유의 이면에 대해서도 유의할 필요가 있다.")
add_paragraph(document, "일부 전립선암 수술 등 세밀한 조작이 절대적인 경우를 제외하고는, 아직까지 담낭 절제술의 표준 치료는 숙련된 전문의가 복강경과 첨단 필수 수술 기구(Advanced essentials tools)를 이용하여 직접 담낭을 절제하는 것이다. 일부 대학병원에서 담낭 절제술에까지 굳이 고가의 로봇 수술을 강요하는 것은, 환자의 치료적 이득보다는 병원 측의 실험 단계 연구 목적이거나 비급여 항목을 통한 수익 창출을 위한 것일 가능성이 매우 높다. 최상의 수술 결과를 좌우하는 것은 값비싼 기계가 아니라, 의사의 풍부한 임상 경험과 정교한 직접 술기이다.", bold_words=True)
add_paragraph(document, "남양주백병원 소화기센터는 숙련된 일반외과 병동 전문의가 최신의 복강경 장비를 십분 활용하여 가장 안전하고 검증된 표준 치료만을 정직하게 제공하고 있다. 건강검진에서 담석을 진단받았다면, 과도한 비용 부담이나 불필요한 시술을 강요받지 않고 믿을 수 있는 전문의와 상의하여 최적의 치료 시기를 결정하는 것이 현명하다.")

# Image
try:
    img_path = r"C:\Users\USER\.gemini\antigravity\brain\1448649a-8b80-4ac3-89f4-28a4ffd9b85c\nano_banana_laparoscopic_1773806761827.png"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(img_path, width=Inches(5.0))
    
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run("(의학적 정밀도와 첨단 기술을 상징하는 '나노 바나나' 모티브의 시각화 아트워크)")
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor(119, 119, 119)
    caption_run.font.name = 'Malgun Gothic'
    caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
except Exception as e:
    print(f"Warning: Could not add image: {e}")

# Apply global font styling to all runs
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        if not run.font.name:
            run.font.name = 'Malgun Gothic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

output_path = r"C:\Users\USER\OneDrive\바탕 화면\언론기사\2026년 3월 언론기사\남양주백병원_담석증_기사.docx"
document.save(output_path)
print(f"Docx saved to {output_path}")
