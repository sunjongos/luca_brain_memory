import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    if level == 0:
        run.font.color.rgb = RGBColor(44, 62, 80)
    else:
        run.font.color.rgb = RGBColor(41, 128, 185)

def add_paragraph(doc, text, bold_words=None):
    p = doc.add_paragraph()
    
    if bold_words:
        # Highlight logic for blog post
        if "급성 담낭염" in text:
            parts = text.split("'급성 담낭염'")
            if len(parts) > 1:
                p.add_run(parts[0])
                run = p.add_run("'급성 담낭염'")
                run.bold = True
                p.add_run(parts[1])
                return p
                
        if "패혈증" in text:
            parts = text.split("'패혈증'")
            if len(parts) > 1:
                p.add_run(parts[0])
                run = p.add_run("'패혈증'")
                run.bold = True
                run.font.color.rgb = RGBColor(217, 83, 79)
                p.add_run(parts[1])
                return p
                
        if "'복강경 수술'" in text:
            parts = text.split("'복강경 수술'")
            if len(parts) > 1:
                p.add_run(parts[0])
                run = p.add_run("'복강경 수술'")
                run.bold = True
                run.font.color.rgb = RGBColor(41, 128, 185)
                p.add_run(parts[1])
                return p
                
        if "수술 후 조기 회복 프로그램(ERAS)" in text:
            parts = text.split("'수술 후 조기 회복 프로그램(ERAS)'")
            if len(parts) > 1:
                p.add_run(parts[0])
                run = p.add_run("'수술 후 조기 회복 프로그램(ERAS)'")
                run.bold = True
                run.font.color.rgb = RGBColor(22, 160, 133)
                p.add_run(parts[1])
                return p
                
        if "집도의의 풍부한 임상 경험과 정교한 손끝 감각" in text:
            parts = text.split("'집도의의 풍부한 임상 경험과 정교한 손끝 감각'")
            if len(parts) > 1:
                p.add_run(parts[0])
                run = p.add_run("'집도의의 풍부한 임상 경험과 정교한 손끝 감각'")
                run.bold = True
                p.add_run(parts[1])
                return p

    p.add_run(text)
    return p

document = Document()

# Title
title = document.add_heading('🏥 [건강정보] 건강검진에서 담석 발견?! 수술, 무조건 해야 할까요? 🩺', 0)
title_run = title.runs[0]
title_run.font.name = 'Malgun Gothic'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# Author
author_p = document.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author_p.add_run('👨‍⚕️ 남양주백병원 소화기센터 김요운 과장(일반외과 전문의)')
author_run.bold = True
author_run.font.name = 'Malgun Gothic'
author_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
author_run.font.color.rgb = RGBColor(22, 160, 133)

# Intro
add_paragraph(document, "안녕하세요! 남양주백병원 소화기센터입니다. 😊", bold_words=True)
add_paragraph(document, "최근 건강검진을 받으시다가 우연히 '담석'이 있다는 결과를 받고 놀라서 병원을 찾으시는 분들이 많습니다. 당장 아프지는 않은데 수술을 해야 할지, 아니면 그냥 지켜봐도 될지 고민이 많으실 텐데요!")
add_paragraph(document, "오늘은 저희 소화기센터 김요운 과장님과 함께 담석증 수술의 모든 것에 대해 속 시원하게 풀어보겠습니다! 💡")

# Sec 1
add_heading(document, "💥 1. 담낭 절제술, 안 아픈데 꼭 해야 하나요?", 2)
add_paragraph(document, '"지금은 안 아프니까 괜찮겠지..." 하고 방치하다가는 큰일 날 수 있습니다! 🚨', bold_words=True)
add_paragraph(document, "담석증을 가볍게 여겨 치료 골든타임을 놓치면 '급성 담낭염'으로 악화될 수 있습니다. 염증이 심해지면 담낭이 터지거나 염증이 온몸으로 퍼지는 '패혈증'을 유발할 수 있어요. 패혈증은 자칫 생명까지 위협할 수 있는 매우 무서운 상태입니다.", bold_words=True)
add_paragraph(document, "따라서 소화불량 같은 가벼운 증상이나 가끔씩 콕콕 찌르는 통증, 혹은 초음파상 염증 소견이 조금이라도 보인다면 예방적 차원에서라도 담낭 절제술을 적극적으로 고려하셔야 합니다.")

# Sec 2
add_heading(document, "✨ 2. 요즘 대세는 '복강경 수술' + '초고속 회복(ERAS)'!", 2)
add_paragraph(document, '"수술하면 크게 째야 하고 아프지 않나요? ㅠㅠ"')
add_paragraph(document, "걱정 마세요! 요즘 담낭 절제술은 대부분 작게 구멍만 내는 '복강경 수술'로 진행됩니다. 상처가 작으니 당연히 통증도 덜하고 회복도 빠르죠. 🏃‍♂️🏃‍♀️", bold_words=True)
add_paragraph(document, "특히 저희 남양주백병원 은 여기서 한 발 더 나아가 '수술 후 조기 회복 프로그램(ERAS)'을 적극적으로 적용하고 있습니다! 환자분들의 통증은 확~ 줄이고, 가장 안전하고 빠르게 일상으로 복귀하실 수 있도록 수술 전후 모든 과정을 체계적으로 관리해 드리고 있답니다. 👍", bold_words=True)

# Sec 3
add_heading(document, "🤖 3. 대학병원 로봇 수술 권유, 무조건 좋은 걸까요?", 2)
add_paragraph(document, '최근 대형 대학병원에서 로봇 수술을 많이 권유받고 오시는 분들이 계십니다. "비싸니까 더 좋겠지?"라고 생각하실 수 있지만, 그 이면을 한 번 살펴볼 필요가 있습니다. 🤔')
add_paragraph(document, "일부 전립선암 수술처럼 현미경 수준의 세밀한 조작이 절대적인 경우를 제외하고는, 담낭 절제술의 '글로벌 표준 치료'는 아직까지도 의사가 직접 첨단 수술 기구(Advanced Essentials Tools)와 복강경을 이용하여 세밀하게 담석을 제거하는 것입니다! 👨‍⚕️✨")
add_paragraph(document, "대학병원에서 굳이 고가의 비급여 로봇 수술을 강요하는 것은, 환자분의 치료 효과 자체보다는 병원의 실험 단계 연구 목적이거나 비급여를 통한 수익 창출을 위한 수단일 가능성도 존재합니다.")
add_paragraph(document, "수술의 성공을 좌우하는 것은 번쩍이는 비싼 기계가 아니라, '집도의의 풍부한 임상 경험과 정교한 손끝 감각'입니다! 💯", bold_words=True)

# Sec 4
add_heading(document, "🏥 믿고 맡길 수 있는 남양주백병원 소화기센터", 2)
add_paragraph(document, "저희 남양주백병원 소화기센터는 무리하게 고가의 비급여 수술을 권유하지 않습니다. 풍부한 경험을 갖춘 병동 전담 전문의가 최신 복강경 장비를 십분 활용하여, 가장 안전하고 검증된 ‘표준 치료’만을 정직하게 제공해 드립니다. 🤝")
add_paragraph(document, "검진에서 담석 진단을 받으셨나요? 더 이상 혼자 고민하지 마시고, 부담 없이 남양주백병원 소화기센터로 오셔서 현명한 치료 계획을 세워보세요! 여러분의 든든한 건강 파트너가 되겠습니다! 💪💚")

# Image
try:
    img_path = r"C:\Users\USER\.gemini\antigravity\brain\1448649a-8b80-4ac3-89f4-28a4ffd9b85c\nano_banana_laparoscopic_1773806761827.png"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(img_path, width=Inches(5.0))
    
    caption_p = document.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run("(초정밀 의학 기술을 상징하는 사이버네틱 '나노 바나나' 아트워크)")
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_run.font.color.rgb = RGBColor(119, 119, 119)
    caption_run.font.name = 'Malgun Gothic'
    caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
except Exception as e:
    print(f"Warning: Could not add image: {e}")

# Hashtags
hashtag_p = document.add_paragraph()
hashtag_run = hashtag_p.add_run("#남양주백병원 #소화기센터 #김요운과장 #담석증 #담낭절제술 #복강경수술 #ERAS #단기회복 #로봇수술팩트체크 #건강검진 #남양주외과 #정직한진료")
hashtag_run.font.color.rgb = RGBColor(52, 152, 219)
hashtag_run.font.name = 'Malgun Gothic'
hashtag_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# Apply global font styling to all runs
for paragraph in document.paragraphs:
    for run in paragraph.runs:
        if not run.font.name:
            run.font.name = 'Malgun Gothic'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

output_path = r"C:\Users\USER\OneDrive\바탕 화면\언론기사\2026년 3월 언론기사\남양주백병원_담석증_블로그.docx"
document.save(output_path)
print(f"Blog Docx saved to {output_path}")
